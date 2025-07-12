# import tkinter as tk
# from tkinter import filedialog, messagebox
# import pyttsx3
# import threading
# import PyPDF2
# import docx
# import speech_recognition as sr

# # Initialize TTS engine
# engine = pyttsx3.init()
# engine.setProperty('rate', 150)

# # Speak Thread
# speak_thread = None

# def extract_text(file_path):
#     try:
#         if file_path.endswith(".pdf"):
#             with open(file_path, 'rb') as file:
#                 reader = PyPDF2.PdfReader(file)
#                 return ''.join(page.extract_text() or '' for page in reader.pages)
#         elif file_path.endswith(".docx"):
#             doc = docx.Document(file_path)
#             return '\n'.join(para.text for para in doc.paragraphs)
#         elif file_path.endswith(".txt"):
#             with open(file_path, 'r', encoding='utf-8') as f:
#                 return f.read()
#         else:
#             return None
#     except Exception as e:
#         messagebox.showerror("Error", f"Failed to read file: {e}")
#         return None

# def speak_text(text):
#     engine.say(text)
#     engine.runAndWait()

# def start_speaking():
#     global speak_thread
#     file_path = filedialog.askopenfilename(filetypes=[("Text or Document Files", "*.txt *.pdf *.docx")])
#     if not file_path:
#         return

#     text = extract_text(file_path)
#     if text:
#         speak_thread = threading.Thread(target=speak_text, args=(text,))
#         speak_thread.start()

# def stop_speaking():
#     engine.stop()

# def listen_for_command():
#     recognizer = sr.Recognizer()
#     with sr.Microphone() as source:
#         status_label.config(text="🎙️ Listening for command...")
#         app.update()
#         audio = recognizer.listen(source)
#     try:
#         command = recognizer.recognize_google(audio).lower()
#         status_label.config(text=f"🗣️ You said: {command}")
#         app.update()

#         if "read" in command:
#             start_speaking()
#         elif "stop" in command:
#             stop_speaking()
#         elif "exit" in command or "close" in command:
#             app.quit()
#         else:
#             messagebox.showinfo("Command", f"Unknown command: {command}")
#     except sr.UnknownValueError:
#         status_label.config(text="😕 Could not understand")
#     except sr.RequestError as e:
#         status_label.config(text=f"❌ Speech error: {e}")

# # GUI Setup
# app = tk.Tk()
# app.title("AI Voice Reader")
# app.geometry("420x250")

# tk.Label(app, text="AI Voice Reader", font=("Arial", 16, "bold")).pack(pady=10)

# tk.Button(app, text="📂 Select File & Speak", command=start_speaking, width=25, bg="lightgreen").pack(pady=5)
# tk.Button(app, text="🛑 Stop Reading", command=stop_speaking, width=25, bg="salmon").pack(pady=5)
# tk.Button(app, text="🎙️ Voice Command", command=listen_for_command, width=25, bg="lightblue").pack(pady=10)

# status_label = tk.Label(app, text="", font=("Arial", 10), fg="gray")
# status_label.pack()

# app.mainloop()



import tkinter as tk
from tkinter import filedialog, messagebox
import pyttsx3
import threading
import PyPDF2
import docx
import speech_recognition as sr


# Initialize TTS engine
engine = pyttsx3.init()
engine.setProperty('rate', 150)

# Speak Thread
speak_thread = None

def extract_text(file_path):
    try:
        if file_path.endswith(".pdf"):
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return ''.join(page.extract_text() or '' for page in reader.pages)
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            return '\n'.join(para.text for para in doc.paragraphs)
        elif file_path.endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return None
    except Exception as e:
        messagebox.showerror("Error", f"Failed to read file: {e}")
        return None

def speak_text(text):
    engine.say(text)
    engine.runAndWait()

def start_speaking():
    global speak_thread
    file_path = filedialog.askopenfilename(filetypes=[("Text or Document Files", "*.txt *.pdf *.docx")])
    if not file_path:
        return

    text = extract_text(file_path)
    if text:
        speak_thread = threading.Thread(target=speak_text, args=(text,))
        speak_thread.start()

def stop_speaking():
    engine.stop()

def listen_for_command():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        status_label.config(text="🎙️ Listening for command...")
        app.update()
        audio = recognizer.listen(source)
    try:
        command = recognizer.recognize_google(audio).lower()
        status_label.config(text=f"🗣️ You said: {command}")
        app.update()

        if "read" in command:
            start_speaking()
        elif "stop" in command:
            stop_speaking()
        elif "exit" in command or "close" in command:
            app.quit()
        else:
            messagebox.showinfo("Command", f"Unknown command: {command}")
    except sr.UnknownValueError:
        status_label.config(text="😕 Could not understand")
    except sr.RequestError as e:
        status_label.config(text=f"❌ Speech error: {e}")

# GUI Setup
app = tk.Tk()
app.title("AI Voice Reader")
app.geometry("420x250")

tk.Label(app, text="AI Voice Reader", font=("Arial", 16, "bold")).pack(pady=10)

tk.Button(app, text="📂 Select File & Speak", command=start_speaking, width=25, bg="lightgreen").pack(pady=5)
tk.Button(app, text="🛑 Stop Reading", command=stop_speaking, width=25, bg="salmon").pack(pady=5)
tk.Button(app, text="🎙️ Voice Command", command=listen_for_command, width=25, bg="lightblue").pack(pady=10)

status_label = tk.Label(app, text="", font=("Arial", 10), fg="gray")
status_label.pack()

app.mainloop()
